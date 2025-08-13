"""Procesadores de documentos específicos por tipo"""

import PyPDF2
import pdfplumber
import docx
import pandas as pd
from bs4 import BeautifulSoup
from pathlib import Path
from typing import Dict, List, Tuple
import magic

from .utils import detect_encoding, table_to_markdown

class DocumentTypeProcessor:
    """Clase base para procesadores de tipos de documentos"""
    
    @staticmethod
    def detect_file_type(file_path: Path) -> str:
        """Detecta el tipo de archivo"""
        try:
            mime = magic.Magic(mime=True)
            file_type = mime.from_file(str(file_path))
            
            type_mapping = {
                'pdf': 'pdf',
                'word': 'docx',
                'text': 'txt',
                'markdown': 'md',
                'html': 'html',
                'excel': 'xlsx',
                'csv': 'csv'
            }
            
            for key, value in type_mapping.items():
                if key in file_type.lower() or file_path.suffix.lower() == f'.{value}':
                    return value
            
            return 'unknown'
        except:
            # Fallback a extensión
            return file_path.suffix.lower().strip('.')
    
    @staticmethod
    def extract_from_pdf(file_path: Path) -> Tuple[str, Dict]:
        """Extrae texto de PDF con pdfplumber para mejor manejo de tablas"""
        text_parts = []
        metadata = {'pages': 0, 'has_tables': False}
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extraer texto
                    page_text = page.extract_text() or ""
                    
                    # Extraer tablas
                    tables = page.extract_tables()
                    if tables:
                        metadata['has_tables'] = True
                        for table in tables:
                            # Convertir tabla a markdown
                            table_md = table_to_markdown(table)
                            page_text += f"\n\n{table_md}\n\n"
                    
                    text_parts.append(f"[Página {page_num}]\n{page_text}")
                
        except Exception as e:
            print(f"    Intentando con PyPDF2: {e}")
            # Fallback a PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['pages'] = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text_parts.append(f"[Página {page_num}]\n{page.extract_text()}")
            except Exception as e2:
                print(f"    ❌ Error con ambos métodos: {e2}")
                return "", metadata
        
        return "\n\n".join(text_parts), metadata
    
    @staticmethod
    def extract_from_docx(file_path: Path) -> str:
        """Extrae texto de Word"""
        doc = docx.Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extraer texto de tablas
        for table in doc.tables:
            table_text = DocumentTypeProcessor.extract_table_from_docx(table)
            text_parts.append(table_text)
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def extract_table_from_docx(table) -> str:
        """Extrae tabla de Word como markdown"""
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        
        if not rows:
            return ""
        
        # Convertir a markdown
        return table_to_markdown(rows)
    
    @staticmethod
    def extract_from_spreadsheet(file_path: Path) -> str:
        """Extrae texto de Excel/CSV"""
        try:
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convertir a markdown
            return df.to_markdown()
        except Exception as e:
            print(f"    ❌ Error procesando spreadsheet: {e}")
            return ""